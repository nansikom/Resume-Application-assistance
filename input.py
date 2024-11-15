import os
from docx import Document
from PyPDF2 import PdfReader
from flask import Flask,request,render_template
import pdfplumber
import re

app = Flask(__name__,template_folder='templates')
@app.route('/')
def hello():
    return render_template('index.html')

@app.route('/upload',methods=['POST'])
def upload():
    files = request.files.getlist('files')
    text_contents = []
    for file in files:
        if file.filename  =='':
            return "No selected file",400
        #more like saving the uploaded file into the file storage object file       
        if file:
            upload_dir= "uploads"
            os.makedirs(upload_dir, exist_ok=True)
        if not os.path.exists(upload_dir):
            return "Upload directory not created.", 500  # Server Error

        savepath=os.path.join("uploads", file.filename)
        file.save(savepath)
        file_ext=os.path.splitext(file.filename)[1].lower()
        text_content = ""
        if file_ext == ".pdf":
            with pdfplumber.open(savepath) as pdf:
                #reader= PdfReader(f)
                for page in pdf.pages:
                    page_text = page.extract_text() 
                    print(page_text)
                    if page_text:
                        text_content += page_text + '\n'
                        print(text_content)
        elif file_ext == ".docx":
            doc = Document(savepath)
            for para in doc.paragraphs:
                    text_content += para.text + '\n'
                    print(text_content)
        text_contents.append(text_content)
    return render_template('index.html',extracted_texts=text_contents)
          



@app.route('/uploadform1',methods=['POST'])
def uploadform1():
    files = request.files.getlist('files')
    text_contents = []
    for file in files:
        if file.filename  =='':
            return "No selected file",400
        #more like saving the uploaded file into the file storage object file       
        if file:
            upload_dir= "uploads/form1"
            os.makedirs(upload_dir, exist_ok=True)
        if not os.path.exists(upload_dir):
            return "Upload directory not created.", 500  # Server Error

        savepath=os.path.join("uploads/form1", file.filename)
        file.save(savepath)
        file_ext=os.path.splitext(file.filename)[1].lower()
        text_content = ""
        if file_ext == ".pdf":
            with pdfplumber.open(savepath) as pdf:
                #reader= PdfReader(f)
                for page in pdf.pages:
                    page_text = page.extract_text() 
                    print(page_text)
                    if page_text:
                        text_content += page_text + '\n'
                        print(text_content)
        elif file_ext == ".docx":
            doc = Document(savepath)
            for para in doc.paragraphs:
                    text_content += para.text + '\n'
        text_contents.append(text_content)
        print(text_content)
    return render_template('index.html',extracted_texts1=text_contents, cats= cats)

def wordextraction(text, text2):
        pattern= r'[^\w\s]'
        cleanedsentencestp1= re.sub(pattern,'',text) 
        lowerletters=cleanedsentencestp1.lower()
        tokens=[]
        tokens= lowerletters.split()
        cleanedsentencestp2= re.sub(pattern,'',text2) 
        lowerletter=cleanedsentencestp2.lower()
        token=[]
        token= lowerletter.split()

        for word in tokens:
            if word in token:
                cats=[]
                cats= word
                print("cats drink milk")
        return cats
if __name__== '__main__':
  app.run(debug=True)
  '''
  Find a  way of inputting the documents and searching thru them 
  '''